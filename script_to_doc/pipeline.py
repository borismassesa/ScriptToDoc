import logging
import re
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, Dict, Iterable, List, Optional, Sequence, Tuple

import nltk
from docx import Document
from nltk.tokenize import sent_tokenize

try:
    import spacy
    from spacy.language import Language
except ImportError as exc:  # pragma: no cover - handled at runtime
    raise RuntimeError("spaCy must be installed to use this pipeline") from exc

try:  # Prefer the gensim TextRank implementation when available.
    from gensim.summarization import summarize as _gensim_summarize  # type: ignore

    def text_rank_summarize(text: str, *, ratio: float, word_count: Optional[int], split: bool) -> List[str]:
        return _gensim_summarize(text, ratio=ratio, word_count=word_count, split=split)

except ImportError:
    try:
        from summa.summarizer import summarize as _summa_summarize  # type: ignore

        def text_rank_summarize(text: str, *, ratio: float, word_count: Optional[int], split: bool) -> List[str]:
            kwargs: Dict[str, object] = {"ratio": ratio, "split": split}
            if word_count is not None:
                kwargs["words"] = word_count
            return _summa_summarize(text, **kwargs)

    except ImportError as exc:  # pragma: no cover - handled at runtime
        raise RuntimeError(
            "Neither gensim.summarization nor summa is available; install one to enable summarization."
        ) from exc


LOGGER = logging.getLogger(__name__)

FILLER_WORDS = {"um", "uh", "like", "you know", "sort of", "kind of", "basically", "actually"}


@dataclass
class PipelineConfig:
    summary_ratio: float = 0.2
    summary_word_count: Optional[int] = None
    max_summary_sentences: int = 10
    min_summary_sentences: int = 3
    filler_words: Sequence[str] = field(default_factory=lambda: tuple(sorted(FILLER_WORDS)))
    step_expander: Optional[Callable[[int, str], str]] = None
    document_title: str = "Training Document from Meeting Transcript"
    use_llm_summary: bool = False
    use_llm_expander: bool = False
    use_llm_workflow: bool = False
    llm_summary_model: str = "gpt-4.1-mini"
    llm_expander_model: str = "gpt-4.1-mini"
    llm_workflow_model: str = "gpt-4.1-mini"
    llm_temperature: float = 0.2
    tone: str = "Neutral"
    audience: str = "General"


def ensure_nltk_resources() -> None:
    """Ensure the NLTK resources required for tokenization are available."""
    resources = {
        "punkt": "tokenizers/punkt",
        "punkt_tab": "tokenizers/punkt_tab",
        "stopwords": "corpora/stopwords",
    }
    for resource, location in resources.items():
        try:
            nltk.data.find(location)
        except LookupError:
            nltk.download(resource)


def load_spacy_model(model: str = "en_core_web_sm") -> "Language":
    """Load a spaCy model, downloading it on demand if missing."""
    try:
        return spacy.load(model)
    except OSError:
        from spacy.cli import download

        download(model)
        return spacy.load(model)


def load_transcripts(input_dir: Path) -> List[Tuple[Path, str]]:
    """Load all .txt transcripts within a directory tree."""
    if not input_dir.exists():
        raise FileNotFoundError(f"Input directory not found: {input_dir}")

    transcripts: List[Tuple[Path, str]] = []
    for path in sorted(input_dir.rglob("*.txt")):
        if path.is_file():
            try:
                transcripts.append((path, path.read_text(encoding="utf-8")))
            except UnicodeDecodeError:
                transcripts.append((path, path.read_text(encoding="latin-1")))
    if not transcripts:
        raise FileNotFoundError(f"No transcript files found in {input_dir}")
    return transcripts


def _remove_timestamps(text: str) -> str:
    # Remove timestamps like [00:00:05], (00:00), 00:00:05 -, etc.
    cleaned = re.sub(r"\[?\b\d{1,2}:\d{2}(?::\d{2})?\]?", " ", text)
    cleaned = re.sub(r"\(\d{1,2}:\d{2}(?::\d{2})?\)", " ", cleaned)
    return cleaned


def clean_transcript(raw_text: str, filler_words: Optional[Iterable[str]] = None) -> str:
    """Normalize transcript text by removing timestamps, filler words, and noise."""
    cleaned = _remove_timestamps(raw_text)
    fillers = tuple(filler_words) if filler_words else tuple(FILLER_WORDS)

    for filler in fillers:
        pattern = rf"\b{re.escape(filler)}\b"
        cleaned = re.sub(pattern, "", cleaned, flags=re.IGNORECASE)

    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    sentences = sent_tokenize(cleaned)
    return " ".join(sentence.strip() for sentence in sentences)


def summarize_transcript(
    cleaned_transcript: str,
    *,
    ratio: float,
    word_count: Optional[int],
    max_sentences: int,
    min_sentences: int,
    use_llm: bool,
    llm_model: str,
    llm_temperature: float,
    tone: str,
    audience: str,
) -> Tuple[List[str], Optional[Dict[str, int]]]:
    """Summarize text using TextRank with graceful fallbacks for short input."""
    if not cleaned_transcript:
        return [], None

    sentences = sent_tokenize(cleaned_transcript)
    if len(sentences) <= min_sentences:
        return sentences, None

    usage: Optional[Dict[str, int]] = None

    if use_llm:
        try:
            from .llm import summarize_with_openai

            llm_summary, usage = summarize_with_openai(
                cleaned_transcript,
                model=llm_model,
                max_sentences=max_sentences,
                temperature=llm_temperature,
                tone=tone,
                audience=audience,
            )
            if llm_summary:
                return llm_summary[:max_sentences], usage
        except Exception:
            # Fall back to extractive summarisation if the LLM call fails.
            pass

    try:
        summary_sentences = text_rank_summarize(
            cleaned_transcript,
            ratio=ratio,
            word_count=word_count,
            split=True,
        )
    except (ValueError, IndexError):
        summary_sentences = []

    if not summary_sentences:
        summary_sentences = sentences[:max_sentences]

    return summary_sentences[:max_sentences], usage


def extract_steps(sentences: Sequence[str], nlp: "Language") -> List[str]:
    """Identify likely process steps by looking for imperative sentences."""
    steps: List[str] = []
    for sentence in sentences:
        doc = nlp(sentence.strip())
        if not doc:
            continue
        first_token = doc[0]
        if first_token.pos_ == "VERB" or (first_token.pos_ == "AUX" and len(doc) > 1 and doc[1].pos_ == "VERB"):
            steps.append(sentence.strip())
    return steps


def _default_step_description(idx: int, step: str) -> str:
    return (
        f"Step {idx}: {step}\n\n"
        "Explanation: Provide additional context, best practices, and any required tools or data. "
        "Highlight potential pitfalls and success criteria for the learner."
    )


def expand_steps(
    steps: Sequence[str],
    expander: Optional[Callable[[int, str], str]] = None,
    progress_callback: Optional[Callable[[int, int], None]] = None,
) -> List[Tuple[str, str]]:
    """Create expanded descriptions for each step."""
    expanded: List[Tuple[str, str]] = []
    total = len(steps) if steps else 1
    for idx, step in enumerate(steps, start=1):
        if expander:
            description = expander(idx, step)
        else:
            description = _default_step_description(idx, step)
        expanded.append((f"Step {idx}", description))
        if progress_callback:
            progress_callback(idx, total)
    return expanded


def create_word_document(
    transcript_name: str,
    sections: Sequence[Tuple[str, str]],
    output_dir: Path,
    *,
    title: str,
) -> Path:
    """Build a Word document containing the expanded steps."""
    output_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(title, level=1)

    for heading, body in sections:
        document.add_heading(heading, level=2)
        for paragraph in body.split("\n\n"):
            text = paragraph.strip()
            if not text:
                continue

            if text.startswith("Key Actions:"):
                document.add_paragraph("Key Actions:", style="Heading 3")
                lines = text.splitlines()[1:]
                for line in lines:
                    line = line.strip()
                    if line.startswith("- "):
                        document.add_paragraph(line[2:].strip(), style="List Bullet")
                continue

            if text.startswith("- ") or "\n- " in text:
                for line in text.splitlines():
                    line = line.strip()
                    if line.startswith("- "):
                        document.add_paragraph(line[2:].strip(), style="List Bullet")
                continue

            document.add_paragraph(text)

    output_path = output_dir / f"{transcript_name}_training.docx"
    document.save(output_path)
    return output_path


def process_transcript(
    transcript_path: Path,
    output_dir: Path,
    *,
    config: Optional[PipelineConfig] = None,
    nlp: Optional["Language"] = None,
    progress_callback: Optional[Callable[[str, float], None]] = None,
) -> Dict[str, object]:
    """Process a single transcript file into a Word document."""
    cfg = config or PipelineConfig()
    ensure_nltk_resources()
    nlp_model = nlp or load_spacy_model()

    report = progress_callback or (lambda stage, pct: None)
    start_time = time.perf_counter()
    metrics: Dict[str, object] = {
        "duration_sec": 0.0,
        "tokens": {},
        "tone": cfg.tone,
        "audience": cfg.audience,
    }

    report("load_transcript", 0.05)
    try:
        raw_text = transcript_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw_text = transcript_path.read_text(encoding="latin-1")

    report("clean_transcript", 0.15)
    cleaned = clean_transcript(raw_text, cfg.filler_words)

    summary_sentences: List[str] = []
    plan_data: List[Dict[str, object]] = []

    if cfg.use_llm_workflow:
        try:
            from .llm import generate_training_plan_with_openai

            report("llm_plan", 0.5)
            LOGGER.info(
                "Requesting OpenAI training plan",
                extra={
                    "transcript": str(transcript_path),
                    "model": cfg.llm_workflow_model,
                },
            )
            plan, plan_usage = generate_training_plan_with_openai(
                cleaned,
                model=cfg.llm_workflow_model,
                temperature=cfg.llm_temperature,
                max_steps=cfg.max_summary_sentences,
                tone=cfg.tone,
                audience=cfg.audience,
            )
            if plan_usage:
                metrics.setdefault("tokens", {})["plan"] = plan_usage

            if plan:
                LOGGER.info("OpenAI returned %d plan steps", len(plan))
                sections: List[Tuple[str, str]] = []
                step_texts: List[str] = []
                for idx, item in enumerate(plan, start=1):
                    title = item.get("title") or f"Step {idx}"
                    summary = item.get("summary") or ""
                    details = item.get("details") or ""
                    actions = item.get("actions") or []

                    if summary:
                        summary_sentences.append(summary)

                    plan_entry = {
                        "title": title,
                        "summary": summary,
                        "details": details,
                        "actions": actions,
                    }
                    plan_data.append(plan_entry)

                    body_parts: List[str] = []
                    if summary:
                        body_parts.append(f"Overview: {summary}")
                    if details:
                        body_parts.append(details)
                    if actions:
                        action_lines = "\n".join(f"- {action}" for action in actions)
                        if action_lines:
                            body_parts.append(f"Key Actions:\n{action_lines}")

                    heading = f"Step {idx}: {title}"
                    body = "\n\n".join(part for part in body_parts if part) or title
                    sections.append((heading, body))
                    step_texts.append(f"{title}: {details or summary or title}")

                report("create_document", 0.9)
                doc_path = create_word_document(
                    transcript_path.stem,
                    sections,
                    output_dir,
                    title=cfg.document_title,
                )
                report("completed", 1.0)

                metrics_final = {
                    **metrics,
                    "duration_sec": round(time.perf_counter() - start_time, 3),
                }
                LOGGER.info(
                    "Completed transcript processing",
                    extra={
                        "transcript": str(transcript_path),
                        "metrics": metrics_final,
                    },
                )

                return {
                    "transcript": transcript_path,
                    "cleaned_text": cleaned,
                    "summary_sentences": summary_sentences,
                    "steps": step_texts,
                    "document_path": doc_path,
                    "plan": plan_data,
                    "metrics": metrics_final,
                }
        except Exception as exc:
            LOGGER.exception("OpenAI workflow failed: %s", exc)
            report("llm_failed", 0.45)

    summary_sentences, summary_usage = summarize_transcript(
        cleaned,
        ratio=cfg.summary_ratio,
        word_count=cfg.summary_word_count,
        max_sentences=cfg.max_summary_sentences,
        min_sentences=cfg.min_summary_sentences,
        use_llm=cfg.use_llm_summary,
        llm_model=cfg.llm_summary_model,
        llm_temperature=cfg.llm_temperature,
        tone=cfg.tone,
        audience=cfg.audience,
    )
    report("summarize_transcript", 0.45)
    if summary_usage:
        metrics.setdefault("tokens", {})["summary"] = summary_usage

    steps = extract_steps(summary_sentences, nlp_model)
    if not steps:
        steps = summary_sentences[: cfg.max_summary_sentences]
    report("extract_steps", 0.6)

    expander_callable = cfg.step_expander
    expansions_from_llm: Optional[List[str]] = None
    if cfg.use_llm_expander and steps:
        try:
            from .llm import expand_steps_with_openai

            report("expand_steps", 0.65)
            expansions_from_llm, expand_usage = expand_steps_with_openai(
                steps,
                model=cfg.llm_expander_model,
                temperature=cfg.llm_temperature,
                tone=cfg.tone,
                audience=cfg.audience,
            )
            if expand_usage:
                metrics.setdefault("tokens", {})["expander"] = expand_usage
        except Exception:
            expansions_from_llm = None

    if expansions_from_llm is not None and expander_callable is None:

        def expander_callable(idx: int, step: str) -> str:  # type: ignore[assignment]
            if idx - 1 < len(expansions_from_llm):
                return expansions_from_llm[idx - 1]
            return _default_step_description(idx, step)

    expanded = expand_steps(
        steps,
        expander_callable,
        progress_callback=lambda current, total: report(
            "expand_steps",
            0.65 + 0.2 * (current / max(total, 1)),
        ),
    )

    for idx, (heading, body) in enumerate(expanded, start=1):
        title = heading.split(":", 1)[-1].strip() if ":" in heading else heading
        summary = summary_sentences[idx - 1].strip() if idx - 1 < len(summary_sentences) else ""
        details: List[str] = []
        actions: List[str] = []
        for paragraph in body.split("\n\n"):
            text = paragraph.strip()
            if not text:
                continue
            if text.startswith("Key Actions:"):
                lines = text.splitlines()[1:]
                for line in lines:
                    line = line.strip()
                    if line.startswith("- "):
                        actions.append(line[2:].strip())
                continue
            if text.startswith("Overview:"):
                if not summary:
                    summary = text.replace("Overview:", "", 1).strip()
                else:
                    details.append(text.replace("Overview:", "", 1).strip())
            else:
                details.append(text)

        detail_text = " ".join(details).strip()
        plan_data.append(
            {
                "title": title or f"Step {idx}",
                "summary": summary,
                "details": detail_text,
                "actions": actions,
            }
        )

    report("create_document", 0.9)
    doc_path = create_word_document(
        transcript_path.stem,
        expanded,
        output_dir,
        title=cfg.document_title,
    )
    report("completed", 1.0)

    metrics["duration_sec"] = round(time.perf_counter() - start_time, 3)
    LOGGER.info(
        "Completed transcript processing",
        extra={
            "transcript": str(transcript_path),
            "metrics": metrics,
        },
    )

    return {
        "transcript": transcript_path,
        "cleaned_text": cleaned,
        "summary_sentences": summary_sentences,
        "steps": steps,
        "document_path": doc_path,
        "plan": plan_data,
        "metrics": metrics,
    }


def run_pipeline(
    input_dir: Path,
    output_dir: Path,
    *,
    config: Optional[PipelineConfig] = None,
) -> List[Dict[str, object]]:
    """Process every transcript in the input directory tree."""
    transcripts = load_transcripts(input_dir)
    results: List[Dict[str, object]] = []
    nlp_model = load_spacy_model()
    for transcript_path, _ in transcripts:
        results.append(process_transcript(transcript_path, output_dir, config=config, nlp=nlp_model))
    return results
